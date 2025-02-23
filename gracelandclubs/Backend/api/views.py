import os
import sqlite3
import json
import torch
from django.conf import settings
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, BitsAndBytesConfig
from api.models import User, NutritionPlan
from django.db.models import Q
from django.contrib.auth import authenticate, get_user_model
from rest_framework import status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.tokens import RefreshToken
from .serializers import RegisterSerializer, LoginSerializer, UserSerializer
from django.contrib.auth.hashers import make_password
from urllib.parse import unquote_plus
import re
import io
import docx
import PyPDF2

User = get_user_model()  # Use Django's built-in function to get the user model

# Initialize model and tokenizer lazily
class ModelLoader:
    model = None
    tokenizer = None

    @classmethod
    def get_model(cls):
        if cls.model is None or cls.tokenizer is None:
            model_directory = settings.MODEL_DIRECTORY
            try:
                cls.tokenizer = AutoTokenizer.from_pretrained(model_directory)
                cls.model = AutoModelForSeq2SeqLM.from_pretrained(
                    model_directory,
                    device_map='auto'
                )
            except Exception as e:
                raise ImportError(f"Could not load the AI model: {str(e)}")
        return cls.model, cls.tokenizer

# View to handle form analysis
class AnalyzeFormView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file uploaded.'}, status=400)

        file = request.FILES['file']
        file_type = file.content_type
        content = file.read()

        text = self.extract_text(content, file_type)
        prompt = self.generate_prompt(text)

        model, tokenizer = ModelLoader.get_model()
        inputs = tokenizer(prompt, return_tensors="pt").to(model.device)
        
        with torch.no_grad():
            outputs = model.generate(
                inputs["input_ids"],
                max_new_tokens=500,
                temperature=0.7,
                top_p=0.95,
                do_sample=True
            )
        
        response = tokenizer.decode(outputs[0], skip_special_tokens=True)
        
        try:
            analysis = json.loads(response)
        except json.JSONDecodeError:
            analysis = {"issues": ["Error parsing model response"], "recommendations": ["Please try again"]}
        
        return JsonResponse(analysis)

    @staticmethod
    def extract_text(file_content: bytes, file_type: str) -> str:
        if file_type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            return " ".join(page.extract_text() for page in pdf_reader.pages)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(io.BytesIO(file_content))
            return " ".join(paragraph.text for paragraph in doc.paragraphs)
        return file_content.decode('utf-8')

    @staticmethod
    def generate_prompt(form_content: str) -> str:
        requirements = """
        Please analyze this appropriation form against these requirements:
        1. Must specify if funding is one-time or recurring
        2. All form fields must be filled out
        3. Must show alternative funding sources
        4. Must benefit a large number of people
        5. No transportation or gas money reimbursements allowed
        6. Must have specific date for spending
        7. Must include campus advertising plan
        For each requirement, indicate if it is met or not met. If not met, explain what needs to be added.
        Format your response as a JSON object with 'issues' and 'recommendations' arrays.
        """
        return f"<s>[INST] Here is an appropriation form content: {form_content} {requirements}[/INST]"

User = get_user_model()  # ✅ CORRECT

# --- Model Setup ---
# Set up BitsAndBytesConfig for 8-bit quantization
bnb_config = BitsAndBytesConfig(
    load_in_8bit=True
)

# Define model directory
save_directory = "./flan-t5-local"

try:
    if not os.path.exists(save_directory):
        print("Downloading FLAN-T5 model...")
        tokenizer = AutoTokenizer.from_pretrained("google/flan-t5-small")
        model = AutoModelForSeq2SeqLM.from_pretrained(
            "google/flan-t5-small",
            quantization_config=bnb_config,
            device_map="auto"
        )
        tokenizer.save_pretrained(save_directory)
        model.save_pretrained(save_directory)
        print("Model saved locally.")
    else:
        print("Loading FLAN-T5 from local storage...")
        tokenizer = AutoTokenizer.from_pretrained(save_directory)
        model = AutoModelForSeq2SeqLM.from_pretrained(
            save_directory,
            device_map="auto"
        )
except Exception as e:
    print(f"Error loading model: {e}")
    model = None
    tokenizer = None
class UserDetailView(APIView):
    permission_classes = [IsAuthenticated]  # ✅ Require authentication

    def get(self, request):
        user = request.user  # Get logged-in user
        serializer = UserSerializer(user)
        return Response(serializer.data)
# --- Helper Functions ---
def fetch_data_from_sql(query, db_path):
    """Execute a SQL query and return results."""
    try:
        if not os.path.exists(db_path):
            return {"error": "Database file not found."}

        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute(query)
        rows = cursor.fetchall()
        columns = [description[0] for description in cursor.description]
        conn.close()

        return {"columns": columns, "data": rows}

    except sqlite3.Error as e:
        return {"error": f"SQL Error: {str(e)}"}
    except Exception as e:
        return {"error": f"General Error: {str(e)}"}

# --- Django Views ---
def get_users(request):
    """Retrieve all users"""
    users = User.objects.all().values()
    return JsonResponse(list(users), safe=False)

def get_user_nutrition(request, user_id):
    """Retrieve specific user's nutrition plan"""
    try:
        user = User.objects.get(id=user_id)
        response_data = {
            "name": user.name,
            "age": user.age,
            "weight": user.weight,
            "height": user.height,
            "activity_level": user.activity_level
        }
        return JsonResponse(response_data, safe=False)
    except User.DoesNotExist:
        return JsonResponse({"error": "User not found"}, status=404)

def get_nutrition_data(request):
    """Retrieve all nutrition data"""
    data = NutritionPlan.objects.all().values()
    return JsonResponse(list(data), safe=False)


def search_nutrition_plans(request):
    """Search nutrition plans, ensuring direct matches and handling range queries stored in the database."""
    query_params = request.GET.dict()  # Convert query params to dictionary
    filters = Q()  # Initialize query filter object

    # Get valid model fields
    model_fields = {field.name: field.get_internal_type() for field in NutritionPlan._meta.fields}

    for key, value in query_params.items():
        value = unquote_plus(value).strip()  # Decode URL-encoded values and trim spaces

        if key in model_fields:
            field_type = model_fields[key]

            # Fetch all stored values for this field (including ranges & exact values)
            stored_values = NutritionPlan.objects.values_list(key, flat=True)
            matching_values = set()  # Use a set to prevent duplicates

            # Handle Numeric Fields with Ranges (e.g., "age=20" when database has "18-25")
            if field_type in ["IntegerField", "FloatField"] or key in ["age", "weight", "height"]:
                try:
                    search_value = float(value)  # Convert input value to float

                    for stored_value in stored_values:
                        stored_value_str = str(stored_value).strip()  # Ensure it's a string

                        # Try exact numeric match (works if stored as "20" or "20.0")
                        try:
                            if float(stored_value_str) == search_value:
                                matching_values.add(stored_value_str)
                        except ValueError:
                            pass  # Not a plain numeric value

                        # Use re.search to find a range anywhere in the string (e.g., "18-25" or "18-25 yrs")
                        match = re.search(r"(\d+(?:\.\d+)?)\s*-\s*(\d+(?:\.\d+)?)", stored_value_str)
                        if match:
                            min_val, max_val = map(float, match.groups())
                            if min_val <= search_value <= max_val:
                                matching_values.add(stored_value_str)

                    # Always include "A" to indicate universal match
                    matching_values.add("A")

                    # Apply the filter if we found any matches
                    if matching_values:
                        filters &= Q(**{f"{key}__in": list(matching_values)})

                except ValueError:
                    pass  # Ignore if the search value isn't a valid number

            # Handle String Fields (Case-Insensitive Exact Match)
            elif field_type in ["CharField", "TextField"]:
                filters &= Q(**{f"{key}__iexact": value}) | Q(**{f"{key}": "A"})

    # Ensure filters are valid; if no conditions were added, we use an empty Q() (which returns all)
    if not filters.children:
        filters = Q()

    # Apply filters dynamically
    plans = NutritionPlan.objects.filter(filters).values()  # Return dicts directly

    # Debug output
    print("Query Params:", query_params)
    print("Generated Query:", filters)

    return JsonResponse({"results": list(plans)})


# Helper function to extract text from various file formats
def extract_text(file_content: bytes, file_type: str) -> str:
    if file_type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = " ".join(page.extract_text() for page in pdf_reader.pages)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(io.BytesIO(file_content))
        text = " ".join(paragraph.text for paragraph in doc.paragraphs)
    else:
        text = file_content.decode('utf-8')
    return text

# Helper function to generate the prompt for the AI model
def generate_prompt(form_content: str) -> str:
    return f"""<s>[INST] Here is an appropriation form content:
{form_content}
Please analyze this appropriation form against these requirements:
1. Must specify if funding is one-time or recurring
2. All form fields must be filled out
3. Must show alternative funding sources
4. Must benefit a large number of people
5. No transportation or gas money reimbursements allowed
6. Must have specific date for spending
7. Must include campus advertising plan
For each requirement, indicate if it is met or not met. If not met, explain what needs to be added.
Format your response as a JSON object with 'issues' and 'recommendations' arrays.[/INST]"""

    
@csrf_exempt
def ask(request):
    """Process a question and generate an AI-based response"""
    if 'file' not in request.FILES:
        return JsonResponse({'error': 'No file uploaded.'}, status=400)

    file = request.FILES['file']
    
    try:
        # Example: Save file to disk or process directly
        with open(file.name, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)

        # Process the file as needed (for example, read from it if it's a text file)
        if file.name.endswith('.txt'):
            with open(file.name, 'r') as file_to_read:
                text_data = file_to_read.read()
                # Process text data, for example, use it with an AI model
                # Assume 'model' and 'tokenizer' are preloaded and ready to use
                input_ids = tokenizer(text_data, return_tensors="pt").input_ids
                outputs = model.generate(input_ids)
                generated_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
                return JsonResponse({"answer": generated_text})

        return JsonResponse({'message': 'File processed successfully.'})

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
    
def execute_sql_query(request):
    """Execute SQL query and return structured JSON"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method. Only POST allowed.'}, status=405)

    try:
        data = json.loads(request.body)
        query = data.get("query")
        db_path = settings.DATABASES['default']['NAME']

        if not query:
            return JsonResponse({"error": "Missing required parameter: 'query'"}, status=400)

        sql_result = fetch_data_from_sql(query, db_path)
        if "error" in sql_result:
            return JsonResponse({"error": sql_result["error"]}, status=400)

        return JsonResponse(sql_result, safe=False)

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)
# ✅ Function to generate JWT tokens
@csrf_exempt
def get_tokens_for_user(user):
    refresh = RefreshToken.for_user(user)
    return {
        'refresh': str(refresh),
        'access': str(refresh.access_token),
    }
class RegisterView(APIView):
    def post(self, request):
        serializer = RegisterSerializer(data=request.data)

        if serializer.is_valid():
            # ✅ Ensure password is hashed
            user = User.objects.create(
                username=serializer.validated_data["username"],
                email=serializer.validated_data["email"],
                password=make_password(serializer.validated_data["password"]),  # Hash the password
                age=request.data.get("age"),
                weight=request.data.get("weight"),
                height=request.data.get("height"),
                activity_level=request.data.get("activity_level"),
                vegan=request.data.get("vegan", False),
                vegetarian=request.data.get("vegetarian", False),
                daily_calories=request.data.get("daily_calories", "na"),
            )

            # ✅ Generate JWT token
            token = get_tokens_for_user(user)

            # ✅ Return user data and token
            user_data = UserSerializer(user).data
            return Response({"user": user_data, "token": token}, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
class LoginView(APIView):
    def post(self, request):
        serializer = LoginSerializer(data=request.data)
        if serializer.is_valid():
            username = serializer.validated_data['username']
            password = serializer.validated_data['password']
            user = authenticate(username=username, password=password)
            if user:
                token = get_tokens_for_user(user)
                
                # ✅ Get full user details using serializer
                user_data = UserSerializer(user).data

                return Response({"user": user_data, "token": token}, status=status.HTTP_200_OK)
            return Response({"error": "Invalid credentials"}, status=status.HTTP_401_UNAUTHORIZED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    